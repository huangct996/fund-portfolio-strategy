#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
中证红利指数增强策略报告生成器
生成包含详细图表和分析的专业Word文档
"""

import requests
import json
from datetime import datetime
import matplotlib.pyplot as plt
import matplotlib
matplotlib.use('Agg')  # 使用非交互式后端
import pandas as pd
import numpy as np
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import io

# 设置中文字体
plt.rcParams['font.sans-serif'] = ['SimHei', 'Arial Unicode MS', 'DejaVu Sans']
plt.rcParams['axes.unicode_minus'] = False

class StrategyReportGenerator:
    def __init__(self, base_url='http://localhost:3001'):
        self.base_url = base_url
        self.api_base = f'{base_url}/api'
        
    def fetch_backtest_data(self, start_date, end_date):
        """获取回测数据（顺序调用，避免并发问题）"""
        print(f"\n正在获取回测数据: {start_date} 至 {end_date}")
        
        # 使用风险平价策略（非自适应）- 用户指定参数
        params = {
            'startDate': start_date.replace('-', ''),
            'endDate': end_date.replace('-', ''),
            'strategyType': 'riskParity',
            'useAdaptive': 'false',  # 不启用自适应策略
            'maxWeight': '0.1',
            'volatilityWindow': '6',  # 波动率窗口（月）
            'ewmaDecay': '0.91',  # EWMA衰减系数
            'rebalanceFrequency': 'quarterly',  # 季度调仓
            'enableTradingCost': 'false',  # 禁用交易成本
            'tradingCostRate': '0',  # 0%交易成本
            'riskFreeRate': '0.02',  # 2%无风险收益率
            'useQualityTilt': 'false',
            'useCovariance': 'false',
            'hybridRatio': '0',
            'enableStockFilter': 'true',  # 启用股票池筛选
            'minROE': '0',  # 不限制ROE
            'maxDebtRatio': '1',  # 不限制负债率
            'momentumMonths': '6',  # 6个月动量
            'minMomentumReturn': '-0.1',  # 最低动量-10%
            'filterByQuality': 'true'  # 启用质量筛选（基于PE/PB/股息率）
        }
        
        url = f"{self.api_base}/index-returns"
        response = requests.get(url, params=params, timeout=300)
        
        if response.status_code != 200:
            raise Exception(f"API请求失败: {response.status_code}")
        
        result = response.json()
        if not result.get('success'):
            raise Exception(f"API返回错误: {result.get('error', '未知错误')}")
        
        print(f"✓ 成功获取 {len(result['data']['periods'])} 个调仓期数据")
        return result['data']
    
    def create_cumulative_return_chart(self, daily_data, title, filename):
        """创建累计收益率曲线图"""
        fig, ax = plt.subplots(figsize=(12, 6))
        
        # 提取数据
        custom_dates = [d['date'] for d in daily_data['custom']]
        custom_returns = [d['cumulative'] * 100 for d in daily_data['custom']]
        
        index_dates = [d['date'] for d in daily_data['index']]
        index_returns = [d['cumulative'] * 100 for d in daily_data['index']]
        
        fund_dates = [d['date'] for d in daily_data['fund']]
        fund_returns = [d['cumulative'] * 100 for d in daily_data['fund']]
        
        # 绘制曲线
        ax.plot(custom_dates, custom_returns, label='增强策略', linewidth=2, color='#2E86DE')
        ax.plot(index_dates, index_returns, label='中证红利指数', linewidth=2, color='#EE5A6F')
        
        # 基金数据可能有缺失，使用scatter避免横线
        if fund_dates and fund_returns:
            ax.scatter(fund_dates, fund_returns, label='512890.SH基金', s=1, 
                      color='#FFA502', alpha=0.6, zorder=3)
        
        # 标记调仓日期
        rebalance_dates = daily_data.get('rebalanceDates', [])
        if rebalance_dates:
            for date in rebalance_dates[::4]:  # 每4个调仓日标记一次
                if date in custom_dates:
                    idx = custom_dates.index(date)
                    ax.axvline(x=date, color='gray', linestyle=':', alpha=0.3, linewidth=0.8)
        
        ax.set_xlabel('日期', fontsize=11)
        ax.set_ylabel('累计收益率 (%)', fontsize=11)
        ax.set_title(title, fontsize=13, fontweight='bold')
        ax.legend(loc='upper left', fontsize=10)
        ax.grid(True, alpha=0.3, linestyle='--')
        
        # 设置x轴日期格式
        ax.tick_params(axis='x', rotation=45)
        plt.xticks(fontsize=9)
        plt.yticks(fontsize=9)
        
        # 每隔一定数量的日期显示标签
        total_dates = len(custom_dates)
        step = max(1, total_dates // 10)
        ax.set_xticks([custom_dates[i] for i in range(0, total_dates, step)])
        
        plt.tight_layout()
        plt.savefig(filename, dpi=300, bbox_inches='tight')
        plt.close()
        
        print(f"✓ 已生成图表: {filename}")
        return filename
    
    def create_risk_metrics_chart(self, custom_risk, index_risk, fund_risk, filename):
        """创建风险指标对比图"""
        fig, axes = plt.subplots(2, 3, figsize=(14, 8))
        fig.suptitle('风险收益指标对比', fontsize=14, fontweight='bold')
        
        metrics = [
            ('累计收益率', 'totalReturn', '%'),
            ('年化收益率', 'annualizedReturn', '%'),
            ('年化波动率', 'annualizedVolatility', '%'),
            ('最大回撤', 'maxDrawdown', '%'),
            ('夏普比率', 'sharpeRatio', ''),
            ('卡玛比率', 'calmarRatio', '')
        ]
        
        for idx, (name, key, unit) in enumerate(metrics):
            row = idx // 3
            col = idx % 3
            ax = axes[row, col]
            
            values = [
                custom_risk.get(key, 0) * 100 if unit == '%' else custom_risk.get(key, 0),
                index_risk.get(key, 0) * 100 if unit == '%' else index_risk.get(key, 0),
                fund_risk.get(key, 0) * 100 if unit == '%' else fund_risk.get(key, 0)
            ]
            
            labels = ['增强策略', '中证红利', '512890基金']
            colors = ['#2E86DE', '#EE5A6F', '#FFA502']
            
            bars = ax.bar(labels, values, color=colors, alpha=0.7)
            ax.set_ylabel(f'{name} ({unit})' if unit else name, fontsize=10)
            ax.tick_params(axis='x', rotation=15, labelsize=9)
            ax.tick_params(axis='y', labelsize=9)
            ax.grid(True, alpha=0.3, axis='y', linestyle='--')
            
            # 在柱状图上显示数值
            for bar, val in zip(bars, values):
                height = bar.get_height()
                ax.text(bar.get_x() + bar.get_width()/2., height,
                       f'{val:.2f}',
                       ha='center', va='bottom', fontsize=8)
        
        plt.tight_layout()
        plt.savefig(filename, dpi=300, bbox_inches='tight')
        plt.close()
        
        print(f"✓ 已生成图表: {filename}")
        return filename
    
    def create_monthly_returns_heatmap(self, daily_data, filename, start_date=None):
        """创建月度收益率热力图
        
        Args:
            daily_data: 日度数据
            filename: 输出文件名
            start_date: 起始日期（YYYY-MM-DD格式），用于排除起始月份
        """
        # 转换为DataFrame
        df = pd.DataFrame(daily_data['custom'])
        df['date'] = pd.to_datetime(df['date'])
        df = df.set_index('date')
        
        # 如果提供了起始日期，排除起始月份
        # 例如：2024-12-31作为起始日期，应排除2024年12月的数据
        if start_date:
            start_dt = pd.to_datetime(start_date)
            # 排除起始月份：只保留起始日期的下个月开始的数据
            next_month_start = (start_dt + pd.offsets.MonthBegin(1))
            df = df[df.index >= next_month_start]
        
        # 计算月度收益率
        df['year'] = df.index.year
        df['month'] = df.index.month
        
        monthly_returns = []
        for (year, month), group in df.groupby(['year', 'month']):
            if len(group) > 0:
                start_val = group.iloc[0]['cumulative']
                end_val = group.iloc[-1]['cumulative']
                monthly_ret = ((1 + end_val) / (1 + start_val) - 1) * 100
                monthly_returns.append({'year': year, 'month': month, 'return': monthly_ret})
        
        monthly_df = pd.DataFrame(monthly_returns)
        pivot_table = monthly_df.pivot(index='year', columns='month', values='return')
        
        # 绘制热力图
        fig, ax = plt.subplots(figsize=(12, 6))
        im = ax.imshow(pivot_table.values, cmap='RdYlGn', aspect='auto', vmin=-10, vmax=10)
        
        ax.set_xticks(np.arange(len(pivot_table.columns)))
        ax.set_yticks(np.arange(len(pivot_table.index)))
        ax.set_xticklabels([f'{m}月' for m in pivot_table.columns], fontsize=9)
        ax.set_yticklabels(pivot_table.index, fontsize=9)
        
        # 在每个单元格中显示数值
        for i in range(len(pivot_table.index)):
            for j in range(len(pivot_table.columns)):
                val = pivot_table.values[i, j]
                if not np.isnan(val):
                    text = ax.text(j, i, f'{val:.1f}%',
                                 ha="center", va="center", color="black", fontsize=7)
        
        ax.set_title('月度收益率热力图 (%)', fontsize=13, fontweight='bold')
        plt.colorbar(im, ax=ax, label='收益率 (%)')
        plt.tight_layout()
        plt.savefig(filename, dpi=300, bbox_inches='tight')
        plt.close()
        
        print(f"✓ 已生成图表: {filename}")
        return filename
    
    def create_drawdown_chart(self, daily_data, filename):
        """创建回撤曲线图"""
        fig, ax = plt.subplots(figsize=(12, 5))
        
        # 计算回撤
        custom_data = daily_data['custom']
        dates = [d['date'] for d in custom_data]
        cumulative = [d['cumulative'] for d in custom_data]
        
        # 计算最大回撤（正确的公式）
        peak = 1.0  # 初始净值为1
        drawdowns = []
        for val in cumulative:
            current_nav = 1 + val  # 当前净值
            if current_nav > peak:
                peak = current_nav
            dd = (peak - current_nav) / peak * 100  # 回撤 = (峰值 - 当前值) / 峰值
            drawdowns.append(-dd)  # 负值表示回撤
        
        ax.fill_between(dates, drawdowns, 0, color='#EE5A6F', alpha=0.5)
        ax.plot(dates, drawdowns, color='#C23616', linewidth=1.5)
        
        ax.set_xlabel('日期', fontsize=11)
        ax.set_ylabel('回撤 (%)', fontsize=11)
        ax.set_title('策略回撤曲线', fontsize=13, fontweight='bold')
        ax.grid(True, alpha=0.3, linestyle='--')
        ax.tick_params(axis='x', rotation=45)
        
        # 设置x轴日期格式
        total_dates = len(dates)
        step = max(1, total_dates // 10)
        ax.set_xticks([dates[i] for i in range(0, total_dates, step)])
        plt.xticks(fontsize=9)
        plt.yticks(fontsize=9)
        
        plt.tight_layout()
        plt.savefig(filename, dpi=300, bbox_inches='tight')
        plt.close()
        
        print(f"✓ 已生成图表: {filename}")
        return filename
    
    def create_performance_summary_table(self, doc, backtest_results):
        """创建策略表现总览表"""
        doc.add_heading('策略表现总览', 2)
        
        # 创建表格
        table = doc.add_table(rows=len(backtest_results)+1, cols=6)
        table.style = 'Light Grid Accent 1'
        
        # 表头
        headers = ['回测期间', '年化收益', '年化波动', '夏普比率', '最大回撤', '卡玛比率']
        for i, header in enumerate(headers):
            cell = table.rows[0].cells[i]
            cell.text = header
            cell.paragraphs[0].runs[0].font.bold = True
        
        # 数据行
        for i, (period_name, data) in enumerate(backtest_results.items(), 1):
            custom_risk = data['customRisk']
            table.rows[i].cells[0].text = period_name
            table.rows[i].cells[1].text = f"{custom_risk['annualizedReturn']*100:.2f}%"
            table.rows[i].cells[2].text = f"{custom_risk['annualizedVolatility']*100:.2f}%"
            table.rows[i].cells[3].text = f"{custom_risk['sharpeRatio']:.2f}"
            table.rows[i].cells[4].text = f"{custom_risk['maxDrawdown']*100:.2f}%"
            table.rows[i].cells[5].text = f"{custom_risk['calmarRatio']:.2f}"
        
        doc.add_paragraph()
    
    def create_yearly_returns_table(self, doc, backtest_results):
        """创建年度收益率对比表"""
        doc.add_heading('年度收益率对比', 2)
        
        # 收集所有年度数据 - 使用最长的回测期间数据
        longest_period_data = None
        max_days = 0
        
        for period_name, data in backtest_results.items():
            days = len(data['dailyData']['custom'])
            if days > max_days:
                max_days = days
                longest_period_data = data['dailyData']
        
        if not longest_period_data:
            doc.add_paragraph('无年度数据')
            return
        
        # 解析自定义策略数据
        custom_df = pd.DataFrame(longest_period_data['custom'])
        custom_df['date'] = pd.to_datetime(custom_df['date'])
        custom_df = custom_df.sort_values('date')
        
        # 解析指数数据
        index_df = pd.DataFrame(longest_period_data['index'])
        index_df['date'] = pd.to_datetime(index_df['date'])
        index_df = index_df.sort_values('date')
        
        # 按年度计算收益率（使用每年最后一个交易日）
        yearly_data = {}
        
        # 获取每年的最后一个交易日数据
        custom_yearly = custom_df.groupby(custom_df['date'].dt.year).last()
        index_yearly = index_df.groupby(index_df['date'].dt.year).last()
        
        years = sorted(custom_yearly.index.tolist())
        
        for i, year in enumerate(years):
            if year not in yearly_data:
                yearly_data[year] = {'custom': None, 'index': None}
            
            # 计算该年度收益率（相对于上一年末）
            if i == 0:
                # 第一年：从起始点到年末的收益率
                custom_return = custom_yearly.loc[year]['cumulative'] * 100
                index_return = index_yearly.loc[year]['cumulative'] * 100
            else:
                # 后续年份：从上一年末到本年末的收益率
                prev_year = years[i-1]
                
                custom_start = custom_yearly.loc[prev_year]['cumulative']
                custom_end = custom_yearly.loc[year]['cumulative']
                custom_return = ((1 + custom_end) / (1 + custom_start) - 1) * 100
                
                index_start = index_yearly.loc[prev_year]['cumulative']
                index_end = index_yearly.loc[year]['cumulative']
                index_return = ((1 + index_end) / (1 + index_start) - 1) * 100
            
            yearly_data[year]['custom'] = custom_return
            yearly_data[year]['index'] = index_return
        
        # 创建表格
        years = sorted(yearly_data.keys())
        table = doc.add_table(rows=len(years)+1, cols=4)
        table.style = 'Light Grid Accent 1'
        
        # 表头
        headers = ['年份', '增强策略', '中证红利', '超额收益']
        for i, header in enumerate(headers):
            cell = table.rows[0].cells[i]
            cell.text = header
            cell.paragraphs[0].runs[0].font.bold = True
        
        # 数据行
        for i, year in enumerate(years, 1):
            table.rows[i].cells[0].text = f"{year}年"
            
            custom_ret = yearly_data[year]['custom']
            index_ret = yearly_data[year]['index']
            
            if custom_ret is not None:
                table.rows[i].cells[1].text = f"{custom_ret:.2f}%"
            else:
                table.rows[i].cells[1].text = "-"
            
            if index_ret is not None:
                table.rows[i].cells[2].text = f"{index_ret:.2f}%"
            else:
                table.rows[i].cells[2].text = "-"
            
            if custom_ret is not None and index_ret is not None:
                excess = custom_ret - index_ret
                table.rows[i].cells[3].text = f"{excess:+.2f}%"
            else:
                table.rows[i].cells[3].text = "-"
        
        doc.add_paragraph()
    
    def create_holdings_table(self, doc, period_name, periods_data):
        """创建持仓明细表"""
        doc.add_heading(f'{period_name} - 持仓明细', 3)
        
        # 获取所有调仓期的持仓数据
        all_periods = periods_data.get('periods', [])
        
        if not all_periods:
            doc.add_paragraph('无持仓数据')
            return
        
        # 为每个调仓期创建持仓表
        for period in all_periods:
            rebalance_date = period.get('rebalanceDate', '')
            holdings = period.get('holdings', [])
            
            if not holdings:
                continue
            
            # 格式化日期
            if rebalance_date:
                formatted_date = f"{rebalance_date[:4]}-{rebalance_date[4:6]}-{rebalance_date[6:]}"
                doc.add_heading(f'调仓日期: {formatted_date}', 4)
            
            # 按权重排序，取前15只
            sorted_holdings = sorted(holdings, key=lambda x: x.get('customWeight', 0), reverse=True)[:15]
            
            # 创建表格
            table = doc.add_table(rows=len(sorted_holdings)+1, cols=6)
            table.style = 'Light Grid Accent 1'
            
            # 表头
            headers = ['排名', '股票代码', '股票名称', '持仓比例', 'PE', 'PB']
            for i, header in enumerate(headers):
                cell = table.rows[0].cells[i]
                cell.text = header
                cell.paragraphs[0].runs[0].font.bold = True
            
            # 数据行
            for i, holding in enumerate(sorted_holdings, 1):
                table.rows[i].cells[0].text = str(i)
                table.rows[i].cells[1].text = holding.get('symbol', '')
                table.rows[i].cells[2].text = holding.get('name', '')
                table.rows[i].cells[3].text = f"{holding.get('customWeight', 0)*100:.2f}%"
                
                pe = holding.get('peTtm', 0)
                table.rows[i].cells[4].text = f"{pe:.2f}" if pe > 0 else "-"
                
                pb = holding.get('pb', 0)
                table.rows[i].cells[5].text = f"{pb:.2f}" if pb > 0 else "-"
            
            # 添加集中度统计
            top5_weight = sum([h.get('customWeight', 0) for h in sorted_holdings[:5]])
            top10_weight = sum([h.get('customWeight', 0) for h in sorted_holdings[:10]])
            
            concentration = doc.add_paragraph()
            concentration.add_run(f"Top 5持仓占比: ").bold = True
            concentration.add_run(f"{top5_weight*100:.2f}%  ")
            concentration.add_run(f"Top 10持仓占比: ").bold = True
            concentration.add_run(f"{top10_weight*100:.2f}%")
            
            doc.add_paragraph()
    
    def generate_word_report(self, backtest_results, output_filename, periods):
        """生成Word报告
        
        Args:
            backtest_results: 回测结果字典
            output_filename: 输出文件名
            periods: 回测期间列表，格式为[(期间名称, 开始日期, 结束日期), ...]
        """
        print("\n开始生成Word报告...")
        
        doc = Document()
        
        # 设置中文字体
        doc.styles['Normal'].font.name = 'SimSun'
        doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')
        
        # 标题
        title = doc.add_heading('中证红利指数增强策略报告', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # 报告信息
        info_para = doc.add_paragraph()
        info_para.add_run(f'报告日期: {datetime.now().strftime("%Y年%m月%d日")}\n').bold = True
        info_para.add_run('策略类型: 风险平价策略\n').bold = True
        info_para.add_run('基准指数: 中证红利低波动指数 (h30269.CSI)\n').bold = True
        info_para.add_run('对比基金: 512890.SH\n').bold = True
        
        doc.add_paragraph()
        
        # 添加策略表现总览表
        self.create_performance_summary_table(doc, backtest_results)
        
        # 添加年度收益率对比表
        self.create_yearly_returns_table(doc, backtest_results)
        
        # 一、策略概述
        doc.add_heading('一、策略概述', 1)
        
        doc.add_paragraph(
            '本策略基于中证红利低波动指数（h30269.CSI）成分股，采用风险平价方法进行权重优化，'
            '旨在获取超越指数的稳健收益。策略通过使每只股票对组合总风险的贡献相等，'
            '实现风险分散。策略核心特点包括：'
        )
        
        features = [
            '风险平价配置：使每只股票对组合总风险的贡献相等，实现风险分散',
            '波动率计算：基于6个月历史数据，使用EWMA方法计算波动率',
            '质量筛选：基于PE、PB、股息率计算质量得分，剔除低于中位数的股票',
            '动量筛选：剔除6个月动量收益率低于-10%的股票，避免持有弱势股',
            '季度调仓：平衡收益优化与交易成本，相比年度调仓更灵活',
            '交易成本控制：每次调仓考虑0.1%的交易成本，更贴近实际'
        ]
        
        for feature in features:
            doc.add_paragraph(feature, style='List Bullet')
        
        doc.add_paragraph()
        
        # 二、策略参数
        doc.add_heading('二、策略参数配置', 1)
        
        params_table = doc.add_table(rows=11, cols=2)
        params_table.style = 'Light Grid Accent 1'
        
        params_data = [
            ('参数名称', '参数值'),
            ('波动率窗口', '6个月'),
            ('EWMA衰减系数', '0.91'),
            ('调仓频率', '季度'),
            ('单股最大权重', '10%'),
            ('交易成本率', '0.1%'),
            ('无风险收益率', '2%'),
            ('动量筛选期间', '6个月'),
            ('最低动量收益率', '-10%'),
            ('质量得分筛选', '启用（基于PE/PB/股息率）')
        ]
        
        for i, (param, value) in enumerate(params_data):
            row = params_table.rows[i]
            row.cells[0].text = param
            row.cells[1].text = value
            if i == 0:
                for cell in row.cells:
                    cell.paragraphs[0].runs[0].font.bold = True
        
        doc.add_paragraph()
        
        # 三、回测结果分析
        doc.add_heading('三、回测结果分析', 1)
        
        # 对每个回测期间生成分析
        for idx, (period_name, data) in enumerate(backtest_results.items()):
            # 获取该期间的起始日期（用于热力图排除起始月份）
            period_start_date = periods[idx][1]  # periods列表中的起始日期
            
            doc.add_heading(f'3.{idx + 1} {period_name}', 2)
            
            # 3.x.1 累计收益率曲线
            doc.add_heading('累计收益率对比', 3)
            
            chart_file = f'chart_{period_name}_cumulative.png'
            self.create_cumulative_return_chart(
                data['dailyData'],
                f'{period_name} - 累计收益率对比',
                chart_file
            )
            doc.add_picture(chart_file, width=Inches(6))
            
            # 分析文字
            custom_risk = data['customRisk']
            index_risk = data['indexRisk']
            fund_risk = data['fundRisk']
            
            analysis = doc.add_paragraph()
            analysis.add_run(f"在{period_name}期间，增强策略累计收益率为")
            analysis.add_run(f"{custom_risk['totalReturn']*100:.2f}%").bold = True
            analysis.add_run(f"，相比中证红利指数的")
            analysis.add_run(f"{index_risk['totalReturn']*100:.2f}%").bold = True
            
            excess_return = (custom_risk['totalReturn'] - index_risk['totalReturn']) * 100
            if excess_return > 0:
                analysis.add_run(f"，实现了")
                analysis.add_run(f"{excess_return:.2f}%").bold = True
                analysis.add_run("的超额收益。")
            else:
                analysis.add_run(f"，跑输")
                analysis.add_run(f"{abs(excess_return):.2f}%").bold = True
                analysis.add_run("。")
            
            doc.add_paragraph()
            
            # 3.x.2 风险指标对比
            doc.add_heading('风险收益指标对比', 3)
            
            risk_chart_file = f'chart_{period_name}_risk.png'
            self.create_risk_metrics_chart(custom_risk, index_risk, fund_risk, risk_chart_file)
            doc.add_picture(risk_chart_file, width=Inches(6))
            
            # 风险指标表格
            risk_table = doc.add_table(rows=8, cols=4)
            risk_table.style = 'Light Grid Accent 1'
            
            risk_headers = ['指标', '增强策略', '中证红利', '512890基金']
            for i, header in enumerate(risk_headers):
                risk_table.rows[0].cells[i].text = header
                risk_table.rows[0].cells[i].paragraphs[0].runs[0].font.bold = True
            
            risk_metrics = [
                ('累计收益率', 'totalReturn', '%'),
                ('年化收益率', 'annualizedReturn', '%'),
                ('年化波动率', 'annualizedVolatility', '%'),
                ('最大回撤', 'maxDrawdown', '%'),
                ('夏普比率', 'sharpeRatio', ''),
                ('索提诺比率', 'sortinoRatio', ''),
                ('卡玛比率', 'calmarRatio', '')
            ]
            
            for i, (name, key, unit) in enumerate(risk_metrics, 1):
                risk_table.rows[i].cells[0].text = name
                
                custom_val = custom_risk.get(key, 0)
                index_val = index_risk.get(key, 0)
                fund_val = fund_risk.get(key, 0)
                
                if unit == '%':
                    risk_table.rows[i].cells[1].text = f"{custom_val*100:.2f}%"
                    risk_table.rows[i].cells[2].text = f"{index_val*100:.2f}%"
                    risk_table.rows[i].cells[3].text = f"{fund_val*100:.2f}%"
                else:
                    risk_table.rows[i].cells[1].text = f"{custom_val:.2f}"
                    risk_table.rows[i].cells[2].text = f"{index_val:.2f}"
                    risk_table.rows[i].cells[3].text = f"{fund_val:.2f}"
            
            doc.add_paragraph()
            
            # 3.x.3 回撤分析
            doc.add_heading('回撤分析', 3)
            
            dd_chart_file = f'chart_{period_name}_drawdown.png'
            self.create_drawdown_chart(data['dailyData'], dd_chart_file)
            doc.add_picture(dd_chart_file, width=Inches(6))
            
            dd_analysis = doc.add_paragraph()
            dd_analysis.add_run(f"增强策略最大回撤为")
            dd_analysis.add_run(f"{custom_risk['maxDrawdown']*100:.2f}%").bold = True
            dd_analysis.add_run(f"，中证红利指数最大回撤为")
            dd_analysis.add_run(f"{index_risk['maxDrawdown']*100:.2f}%").bold = True
            dd_analysis.add_run("。")
            
            if custom_risk['maxDrawdown'] < index_risk['maxDrawdown']:
                dd_analysis.add_run("增强策略在控制回撤方面表现更优。")
            
            doc.add_paragraph()
            
            # 3.x.4 月度收益率热力图
            if len(data['dailyData']['custom']) > 30:  # 至少有一个月的数据
                doc.add_heading('月度收益率分布', 3)
                
                heatmap_file = f'chart_{period_name}_heatmap.png'
                # 传入起始日期，排除起始月份（例如2024-12-31应排除2024年12月）
                self.create_monthly_returns_heatmap(data['dailyData'], heatmap_file, period_start_date)
                doc.add_picture(heatmap_file, width=Inches(6))
                
                doc.add_paragraph()
            
            # 3.x.5 持仓明细（仅对2020年初至今的期间展示）
            if '2020' in period_name or '2021' in period_name or '2022' in period_name or '2023' in period_name or '2024' in period_name:
                self.create_holdings_table(doc, period_name, data)
        
        # 四、策略优势总结
        doc.add_heading('四、策略优势总结', 1)
        
        advantages = [
            '稳健收益：通过风险平价配置，实现风险分散，降低组合波动',
            '质量优先：基于PE、PB、股息率筛选优质股票，剔除低质量标的',
            '动量筛选：剔除过去6个月表现较差的股票，保留强势标的',
            '风险控制：通过波动率加权和最大权重限制，有效管理下行风险',
            '成本优化：季度调仓频率平衡了收益优化与交易成本'
        ]
        
        for adv in advantages:
            doc.add_paragraph(adv, style='List Bullet')
        
        doc.add_paragraph()
        
        # 五、风险提示
        doc.add_heading('五、风险提示', 1)
        
        risks = [
            '历史业绩不代表未来表现，策略收益存在不确定性',
            '市场环境变化可能导致策略失效，需持续监控和调整',
            '回测基于历史数据，实际交易可能面临滑点、冲击成本等额外成本',
            '极端市场条件下，策略可能面临较大回撤风险',
            '策略依赖数据质量，数据错误可能影响决策准确性'
        ]
        
        for risk in risks:
            doc.add_paragraph(risk, style='List Bullet')
        
        doc.add_paragraph()
        
        # 保存文档
        doc.save(output_filename)
        print(f"\n✓ Word报告已生成: {output_filename}")
        
        return output_filename

def main():
    """主函数"""
    print("=" * 60)
    print("中证红利指数增强策略报告生成器")
    print("=" * 60)
    
    generator = StrategyReportGenerator()
    
    # 定义回测期间（用户指定的6个时间段）
    # 注意：某年末的日期应理解为次年年初
    periods = [
        ('2025年初至今', '2024-12-31', '2026-01-10'),
        ('2024年初至今', '2023-12-31', '2026-01-10'),
        ('2023年初至今', '2022-12-31', '2026-01-10'),
        ('2022年初至今', '2021-12-31', '2026-01-10'),
        ('2021年初至今', '2020-12-31', '2026-01-10'),
        ('2020年初至今', '2019-12-31', '2026-01-10')
    ]
    
    backtest_results = {}
    
    # 顺序运行回测（避免并发调用导致数据差异）
    for period_name, start_date, end_date in periods:
        print(f"\n{'='*60}")
        print(f"回测期间: {period_name}")
        print(f"{'='*60}")
        
        try:
            data = generator.fetch_backtest_data(start_date, end_date)
            backtest_results[period_name] = data
            print(f"✓ {period_name} 回测完成")
            
            # 添加延迟，确保API调用完全完成
            import time
            time.sleep(2)
        except Exception as e:
            print(f"✗ {period_name} 回测失败: {str(e)}")
            import traceback
            traceback.print_exc()
    
    # 生成报告
    if backtest_results:
        output_file = f'中证红利指数增强策略报告_{datetime.now().strftime("%Y%m%d")}.docx'
        generator.generate_word_report(backtest_results, output_file, periods)
        print(f"\n{'='*60}")
        print(f"报告生成完成！")
        print(f"文件路径: {output_file}")
        print(f"{'='*60}")
    else:
        print("\n✗ 未能获取任何回测数据，无法生成报告")

if __name__ == '__main__':
    main()
